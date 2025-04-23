import os
import sys
from pathlib import Path

# Proje kök dizinini ayarla
ROOT_DIR = Path(__file__).resolve().parent
sys.path.append(str(ROOT_DIR))

# Gerekli dizinleri oluştur
os.makedirs(ROOT_DIR / "templates", exist_ok=True)
os.makedirs(ROOT_DIR / "output", exist_ok=True)
os.makedirs(ROOT_DIR / "static", exist_ok=True)

# Şablon dosyalarını oluştur
from docx import Document

def create_template_files():
    """Örnek şablon dosyalarını oluştur"""
    # Tüketici şikayet dilekçesi şablonu
    tuketici_template = """
    T.C.
    {{muhatap_ad}}
    {{muhatap_adres}}
    
    Şikayet Eden: {{ad_soyad}}
    T.C. Kimlik No: {{tc_kimlik}}
    Adres: {{adres}}
    Telefon: {{telefon}}
    {% if eposta %}E-posta: {{eposta}}{% endif %}
    
    Konu: {{sikayet_konusu}} hakkında şikayet
    
    Sayın Yetkili,
    
    {{satin_alma_tarihi}} tarihinde {{muhatap_ad}}'dan/den {{urun_hizmet}} satın aldım. Ancak aşağıda detaylarını açıkladığım sorunlarla karşılaştım:
    
    {{olay_detayi}}
    
    Bu nedenle aşağıdaki taleplerde bulunuyorum:
    {% for talep in talepler %}
    - {{talep}}
    {% endfor %}
    
    Gereğinin yapılmasını arz ederim.
    
    Tarih: {{tarih}}
    
    İmza
    {{ad_soyad}}
    """
    
    # İş akdi fesih itiraz dilekçesi şablonu
    is_akdi_template = """
    T.C.
    ÇALIŞMA VE İŞ KURUMU İL MÜDÜRLÜĞÜ
    {{isyeri_unvan}} İşyeri Arabuluculuk Bürosu
    
    İtiraz Eden: {{ad_soyad}}
    T.C. Kimlik No: {{tc_kimlik}}
    Adres: {{adres}}
    Telefon: {{telefon}}
    {% if eposta %}E-posta: {{eposta}}{% endif %}
    
    Karşı Taraf: {{isyeri_unvan}}
    Adres: {{isyeri_adres}}
    
    Konu: İş akdi feshi işlemine itiraz
    
    Sayın Yetkili,
    
    {{ise_baslama_tarihi}} tarihinde başladığım işimden {{is_akdi_fesih_tarihi}} tarihinde {{fesih_sebebi}} gerekçesiyle iş akdim feshedilmiştir.
    
    Ancak bu fesih işlemi aşağıdaki nedenlerle haksızdır:
    
    {{itiraz_sebebi}}
    
    Bu nedenle aşağıdaki taleplerde bulunuyorum:
    {% for talep in talepler %}
    - {{talep}}
    {% endfor %}
    
    Gereğinin yapılmasını arz ederim.
    
    Tarih: {{tarih}}
    
    İmza
    {{ad_soyad}}
    """
    
    # Kira tespit davası dilekçesi şablonu
    kira_tespit_template = """
    T.C.
    {{mahkeme}} SULH HUKUK MAHKEMESİ SAYIN HAKİMLİĞİNE
    
    Davacı: {{ad_soyad}}
    T.C. Kimlik No: {{tc_kimlik}}
    Adres: {{adres}}
    Telefon: {{telefon}}
    {% if eposta %}E-posta: {{eposta}}{% endif %}
    
    Davalı: {{karsi_taraf_ad}}
    Adres: {{karsi_taraf_adres}}
    
    Konu: Kira bedelinin tespiti talebi
    
    Sayın Hakim,
    
    {{kira_baslangic}} tarihinde başlayan kira sözleşmesi ile {{gayrimenkul_adres}} adresindeki gayrimenkulü kiralamış bulunmaktayım. Mevcut kira bedeli aylık {{mevcut_kira}} TL'dir.
    
    Aşağıda belirttiğim gerekçelerle kira bedelinin aylık {{talep_edilen_kira}} TL olarak tespit edilmesini talep ediyorum:
    
    {{talep_sebebi}}
    
    Gereğinin yapılmasını arz ederim.
    
    Tarih: {{tarih}}
    
    İmza
    {{ad_soyad}}
    """
    
    # Şablon dosyalarını oluştur
    # Tüketici şikayet dilekçesi
    doc1 = Document()
    doc1.add_paragraph(tuketici_template)
    doc1.save(ROOT_DIR / "templates" / "tuketici_sikayet.docx")
    
    # İş akdi fesih itiraz dilekçesi
    doc2 = Document()
    doc2.add_paragraph(is_akdi_template)
    doc2.save(ROOT_DIR / "templates" / "is_akdi_fesih_itiraz.docx")
    
    # Kira tespit davası dilekçesi
    doc3 = Document()
    doc3.add_paragraph(kira_tespit_template)
    doc3.save(ROOT_DIR / "templates" / "kira_tespit.docx")

if __name__ == "__main__":
    create_template_files()
    print("Şablon dosyaları başarıyla oluşturuldu.")

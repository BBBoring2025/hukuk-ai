from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import json
import shutil
import uuid
import datetime
from docxtpl import DocxTemplate
from docx2pdf import convert
import uvicorn

# Modelleri tanımla
class User(BaseModel):
    username: str
    email: str
    full_name: Optional[str] = None
    user_type: str  # "avukat", "kobi", "bireysel"

class PetitionTemplate(BaseModel):
    id: str
    name: str
    category: str
    description: str
    file_path: str
    fields: List[Dict[str, Any]]
    popularity: int = 0

class PetitionData(BaseModel):
    template_id: str
    user_data: Dict[str, Any]
    recipient_data: Dict[str, Any]
    content_data: Dict[str, Any]

class PetitionResponse(BaseModel):
    id: str
    file_path: str
    created_at: str
    template_id: str
    template_name: str

# Uygulama oluştur
app = FastAPI(title="Türkiye Hukuk AI Platformu - Dilekçe Üretici API")

# CORS ayarları
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Gerçek uygulamada belirli domainlere izin verilmeli
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Statik dosyalar için klasör
os.makedirs("static", exist_ok=True)
os.makedirs("templates", exist_ok=True)
os.makedirs("output", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

# Örnek şablonlar
templates = [
    PetitionTemplate(
        id="tuketici-sikayet",
        name="Tüketici Şikayet Dilekçesi",
        category="tüketici",
        description="Satın alınan mal veya hizmetle ilgili şikayetleri ilgili kurumlara iletmek için kullanılır.",
        file_path="templates/tuketici_sikayet.docx",
        fields=[
            {"name": "ad_soyad", "label": "Ad Soyad", "type": "text", "required": True},
            {"name": "tc_kimlik", "label": "T.C. Kimlik No", "type": "text", "required": True},
            {"name": "adres", "label": "Adres", "type": "textarea", "required": True},
            {"name": "telefon", "label": "Telefon", "type": "text", "required": True},
            {"name": "eposta", "label": "E-posta", "type": "email", "required": False},
            {"name": "muhatap_tur", "label": "Muhatap Türü", "type": "select", "options": ["Kurum", "Şirket", "Kişi"], "required": True},
            {"name": "muhatap_ad", "label": "Muhatap Adı", "type": "text", "required": True},
            {"name": "muhatap_adres", "label": "Muhatap Adres", "type": "textarea", "required": True},
            {"name": "urun_hizmet", "label": "Satın Alınan Ürün/Hizmet", "type": "text", "required": True},
            {"name": "satin_alma_tarihi", "label": "Satın Alma Tarihi", "type": "date", "required": True},
            {"name": "sikayet_konusu", "label": "Şikayet Konusu", "type": "text", "required": True},
            {"name": "olay_detayi", "label": "Olay Detayı", "type": "textarea", "required": True},
            {"name": "talepler", "label": "Talepler", "type": "checkbox", "options": ["Ürün değişimi", "Ücret iadesi", "Onarım", "Tazminat"], "required": True},
        ],
        popularity=100
    ),
    PetitionTemplate(
        id="is-akdi-fesih-itiraz",
        name="İş Akdi Fesih İtiraz Dilekçesi",
        category="iş",
        description="İş akdinin haksız feshedilmesi durumunda itiraz için kullanılır.",
        file_path="templates/is_akdi_fesih_itiraz.docx",
        fields=[
            {"name": "ad_soyad", "label": "Ad Soyad", "type": "text", "required": True},
            {"name": "tc_kimlik", "label": "T.C. Kimlik No", "type": "text", "required": True},
            {"name": "adres", "label": "Adres", "type": "textarea", "required": True},
            {"name": "telefon", "label": "Telefon", "type": "text", "required": True},
            {"name": "eposta", "label": "E-posta", "type": "email", "required": False},
            {"name": "isyeri_unvan", "label": "İşyeri Ünvanı", "type": "text", "required": True},
            {"name": "isyeri_adres", "label": "İşyeri Adresi", "type": "textarea", "required": True},
            {"name": "ise_baslama_tarihi", "label": "İşe Başlama Tarihi", "type": "date", "required": True},
            {"name": "is_akdi_fesih_tarihi", "label": "İş Akdi Fesih Tarihi", "type": "date", "required": True},
            {"name": "fesih_sebebi", "label": "Bildirilen Fesih Sebebi", "type": "textarea", "required": True},
            {"name": "itiraz_sebebi", "label": "İtiraz Sebebi", "type": "textarea", "required": True},
            {"name": "talepler", "label": "Talepler", "type": "checkbox", "options": ["İşe iade", "Kıdem tazminatı", "İhbar tazminatı", "Diğer alacaklar"], "required": True},
        ],
        popularity=85
    ),
    PetitionTemplate(
        id="kira-tespit",
        name="Kira Tespit Davası Dilekçesi",
        category="kira",
        description="Kira bedelinin tespiti için mahkemeye başvuru dilekçesi.",
        file_path="templates/kira_tespit.docx",
        fields=[
            {"name": "ad_soyad", "label": "Ad Soyad", "type": "text", "required": True},
            {"name": "tc_kimlik", "label": "T.C. Kimlik No", "type": "text", "required": True},
            {"name": "adres", "label": "Adres", "type": "textarea", "required": True},
            {"name": "telefon", "label": "Telefon", "type": "text", "required": True},
            {"name": "eposta", "label": "E-posta", "type": "email", "required": False},
            {"name": "karsi_taraf_ad", "label": "Karşı Taraf Ad Soyad/Ünvan", "type": "text", "required": True},
            {"name": "karsi_taraf_adres", "label": "Karşı Taraf Adres", "type": "textarea", "required": True},
            {"name": "kira_baslangic", "label": "Kira Sözleşmesi Başlangıç Tarihi", "type": "date", "required": True},
            {"name": "kira_bitis", "label": "Kira Sözleşmesi Bitiş Tarihi", "type": "date", "required": False},
            {"name": "gayrimenkul_adres", "label": "Kiralanan Gayrimenkul Adresi", "type": "textarea", "required": True},
            {"name": "mevcut_kira", "label": "Mevcut Kira Bedeli", "type": "number", "required": True},
            {"name": "talep_edilen_kira", "label": "Talep Edilen Kira Bedeli", "type": "number", "required": True},
            {"name": "talep_sebebi", "label": "Talep Sebebi", "type": "textarea", "required": True},
        ],
        popularity=70
    ),
]

# Örnek oluşturulan dilekçeler
petitions = []

# Şablon oluşturma fonksiyonu
def create_template_files():
    """Örnek şablon dosyalarını oluştur"""
    # Tüketici şikayet dilekçesi şablonu
    tuketici_template = """
    T.C.
    {{muhatap_ad}}
    {{muhatap_adres}}
    
    Şikayet Eden: {{ad_soyad}}
    T.C. Kimlik No: {{tc_kimlik}}
    Adres: {{adres}}
    Telefon: {{telefon}}
    {% if eposta %}E-posta: {{eposta}}{% endif %}
    
    Konu: {{sikayet_konusu}} hakkında şikayet
    
    Sayın Yetkili,
    
    {{satin_alma_tarihi}} tarihinde {{muhatap_ad}}'dan/den {{urun_hizmet}} satın aldım. Ancak aşağıda detaylarını açıkladığım sorunlarla karşılaştım:
    
    {{olay_detayi}}
    
    Bu nedenle aşağıdaki taleplerde bulunuyorum:
    {% for talep in talepler %}
    - {{talep}}
    {% endfor %}
    
    Gereğinin yapılmasını arz ederim.
    
    Tarih: {{tarih}}
    
    İmza
    {{ad_soyad}}
    """
    
    # İş akdi fesih itiraz dilekçesi şablonu
    is_akdi_template = """
    T.C.
    ÇALIŞMA VE İŞ KURUMU İL MÜDÜRLÜĞÜ
    {{isyeri_unvan}} İşyeri Arabuluculuk Bürosu
    
    İtiraz Eden: {{ad_soyad}}
    T.C. Kimlik No: {{tc_kimlik}}
    Adres: {{adres}}
    Telefon: {{telefon}}
    {% if eposta %}E-posta: {{eposta}}{% endif %}
    
    Karşı Taraf: {{isyeri_unvan}}
    Adres: {{isyeri_adres}}
    
    Konu: İş akdi feshi işlemine itiraz
    
    Sayın Yetkili,
    
    {{ise_baslama_tarihi}} tarihinde başladığım işimden {{is_akdi_fesih_tarihi}} tarihinde {{fesih_sebebi}} gerekçesiyle iş akdim feshedilmiştir.
    
    Ancak bu fesih işlemi aşağıdaki nedenlerle haksızdır:
    
    {{itiraz_sebebi}}
    
    Bu nedenle aşağıdaki taleplerde bulunuyorum:
    {% for talep in talepler %}
    - {{talep}}
    {% endfor %}
    
    Gereğinin yapılmasını arz ederim.
    
    Tarih: {{tarih}}
    
    İmza
    {{ad_soyad}}
    """
    
    # Kira tespit davası dilekçesi şablonu
    kira_tespit_template = """
    T.C.
    {{mahkeme}} SULH HUKUK MAHKEMESİ SAYIN HAKİMLİĞİNE
    
    Davacı: {{ad_soyad}}
    T.C. Kimlik No: {{tc_kimlik}}
    Adres: {{adres}}
    Telefon: {{telefon}}
    {% if eposta %}E-posta: {{eposta}}{% endif %}
    
    Davalı: {{karsi_taraf_ad}}
    Adres: {{karsi_taraf_adres}}
    
    Konu: Kira bedelinin tespiti talebi
    
    Sayın Hakim,
    
    {{kira_baslangic}} tarihinde başlayan kira sözleşmesi ile {{gayrimenkul_adres}} adresindeki gayrimenkulü kiralamış bulunmaktayım. Mevcut kira bedeli aylık {{mevcut_kira}} TL'dir.
    
    Aşağıda belirttiğim gerekçelerle kira bedelinin aylık {{talep_edilen_kira}} TL olarak tespit edilmesini talep ediyorum:
    
    {{talep_sebebi}}
    
    Gereğinin yapılmasını arz ederim.
    
    Tarih: {{tarih}}
    
    İmza
    {{ad_soyad}}
    """
    
    # Şablon dosyalarını oluştur
    from docx import Document
    
    # Tüketici şikayet dilekçesi
    doc1 = Document()
    doc1.add_paragraph(tuketici_template)
    doc1.save("templates/tuketici_sikayet.docx")
    
    # İş akdi fesih itiraz dilekçesi
    doc2 = Document()
    doc2.add_paragraph(is_akdi_template)
    doc2.save("templates/is_akdi_fesih_itiraz.docx")
    
    # Kira tespit davası dilekçesi
    doc3 = Document()
    doc3.add_paragraph(kira_tespit_template)
    doc3.save("templates/kira_tespit.docx")

# Uygulama başlangıcında şablon dosyalarını oluştur
create_template_files()

# API Endpoint'leri
@app.get("/")
async def root():
    return {"message": "Türkiye Hukuk AI Platformu - Dilekçe Üretici API"}

@app.get("/templates", response_model=List[PetitionTemplate])
async def get_templates():
    """Tüm dilekçe şablonlarını listele"""
    return templates

@app.get("/templates/{category}", response_model=List[PetitionTemplate])
async def get_templates_by_category(category: str):
    """Belirli bir kategorideki dilekçe şablonlarını listele"""
    return [t for t in templates if t.category == category]

@app.get("/template/{template_id}", response_model=PetitionTemplate)
async def get_template(template_id: str):
    """Belirli bir dilekçe şablonunu getir"""
    for template in templates:
        if template.id == template_id:
            return template
    raise HTTPException(status_code=404, detail="Şablon bulunamadı")

@app.post("/generate", response_model=PetitionResponse)
async def generate_petition(petition_data: PetitionData):
    """Dilekçe oluştur"""
    # Şablonu bul
    template = None
    for t in templates:
        if t.id == petition_data.template_id:
            template = t
            break
    
    if not template:
        raise HTTPException(status_code=404, detail="Şablon bulunamadı")
    
    # Şablon dosyasını yükle
    try:
        doc = DocxTemplate(template.file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Şablon dosyası yüklenemedi: {str(e)}")
    
    # Şablona verileri yerleştir
    context = {
        **petition_data.user_data,
        **petition_data.recipient_data,
        **petition_data.content_data,
        "tarih": datetime.datetime.now().strftime("%d/%m/%Y")
    }
    
    try:
        doc.render(context)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Şablon işlenemedi: {str(e)}")
    
    # Benzersiz dosya adı oluştur
    petition_id = str(uuid.uuid4())
    docx_path = f"output/{petition_id}.docx"
    pdf_path = f"output/{petition_id}.pdf"
    
    # Dosyayı kaydet
    try:
        doc.save(docx_path)
        # PDF'e dönüştür
        convert(docx_path, pdf_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Dosya kaydedilemedi: {str(e)}")
    
    # Oluşturulan dilekçeyi kaydet
    petition = PetitionResponse(
        id=petition_id,
        file_path=pdf_path,
        created_at=datetime.datetime.now().isoformat(),
        template_id=template.id,
        template_name=template.name
    )
    petitions.append(petition)
    
    # Şablonun popülerliğini artır
    template.popularity += 1
    
    return petition

@app.get("/download/{petition_id}")
async def download_petition(petition_id: str, format: str = "pdf"):
    """Oluşturulan dilekçeyi indir"""
    for petition in petitions:
        if petition.id == petition_id:
            if format == "pdf":
                file_path = f"output/{petition_id}.pdf"
            else:
                file_path = f"output/{petition_id}.docx"
            
            if os.path.exists(file_path):
                return FileResponse(
                    path=file_path,
                    filename=f"dilekce_{petition_id}.{format}",
                    media_type=f"application/{'pdf' if format == 'pdf' else 'vnd.openxmlformats-officedocument.wordprocessingml.document'}"
                )
            else:
                raise HTTPException(status_code=404, detail="Dosya bulunamadı")
    
    raise HTTPException(status_code=404, detail="Dilekçe bulunamadı")

@app.get("/petitions", response_model=List[PetitionResponse])
async def get_petitions():
    """Oluşturulan tüm dilekçeleri listele"""
    return petitions

@app.get("/popular-templates", response_model=List[PetitionTemplate])
async def get_popular_templates(limit: int = 5):
    """En popüler dilekçe şablonlarını listele"""
    sorted_templates = sorted(templates, key=lambda x: x.popularity, reverse=True)
    return sorted_templates[:limit]

# Ana uygulama
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
